#Alicia Melotik
#amelotik@nd.edu

import argparse
import json
import docx
import os.path
import sys

#create the word doc and initially write the contents of textFile to it
def createWordDoc(textFile):
    doc = docx.Document()
    with open(textFile, 'r') as f:
        lines = f.read()
    doc.add_paragraph(lines)
    return doc

#add table with the statistics data to the passed document
def insertTable(statData, doc):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for key, value in statData.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

#combine all the document elements into singular doc saved into outputFile
def combineDocParts(textString, statDict, graphPNG, outputFile):
    doc = createWordDoc(textString)
    insertTable(statDict, doc)
    doc.add_picture(graphPNG)
    doc.save(outputFile)

#to test this file, get user arguments and call functions
if __name__ == "__main__":  
    parser = argparse.ArgumentParser()
    parser.add_argument("textFile", help="text file to be read in")
    parser.add_argument("image", help="the image to be read in")
    parser.add_argument("outputFile", help="the name of the doc file to output the plot to")
    #use the parser to get the entered arguments
    args = parser.parse_args()

    if not os.path.isfile(args.textFile):
        print("Error: " + args.textFile + " cannot be found")
        sys.exit(1)
    if not os.path.isfile(args.image):
        print("Error: " + args.image + " cannot be found")
        sys.exit(1)
    if os.path.isfile(args.outputFile):
        print("Warning: " + args.outputFile + " already exists. It will be overwritten.")
        sys.exit(1)

    #dummy dictionary for testing
    statDict = {'period': 0, 'interface': 0, 'num points': 0, 'min': 0, 'max': 0, 
                'mean': 0, 'median': 0, 'std dev': 0, '10th percentile': 0, '90th percentile': 0}
    combineDocParts(args.textFile, statDict, args.image, args.outputFile)

