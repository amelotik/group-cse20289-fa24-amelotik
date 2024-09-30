#Alicia Melotik
#amelotik@nd.edu

import argparse
import json
import docx

def createWordDoc(textFile):
    doc = docx.Document()
    with open(textFile, 'r') as f:
        lines = f.read()
    doc.add_paragraph(lines)
    return doc

def insertTable(statData, doc):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for key, value in statData.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

def combineDocParts(textString, statDict, graphPNG, outputFile):
    doc = createWordDoc(textString)
    insertTable(statDict, doc)
    doc.add_picture(graphPNG)
    doc.save(outputFile)

if __name__ == "__main__":  
    parser = argparse.ArgumentParser()
    parser.add_argument("textFile", help="text file to be read in")
    parser.add_argument("image", help="the image to be read in")
    parser.add_argument("outputFile", help="the name of the doc file to output the plot to")
    #use the parser to get the entered arguments
    args = parser.parse_args()


    statDict = {'period': 0, 'interface': 0, 'num points': 0, 'min': 0, 'max': 0, 
                'mean': 0, 'median': 0, 'std dev': 0, '10th percentile': 0, '90th percentile': 0}
    combineDocParts(args.textFile, statDict, args.image, args.outputFile)

